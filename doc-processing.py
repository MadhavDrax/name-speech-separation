import os
from flask import Flask, request, render_template, send_file
from docx import Document
import re

# Set up Flask and point the template folder to the current directory
app = Flask(__name__, template_folder=os.path.dirname(os.path.abspath(__file__)))
UPLOAD_FOLDER = 'public'
OUTPUT_FOLDER = 'public/output'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# Ensure upload and output folders exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Step 1: Parse the DOCX file and extract names and speeches
def parse_doc(input_file):
    doc = Document(input_file)
    data = []
    current_name = ""
    current_speech = []

    # Regular expression to detect names with a pattern like #10#
    name_pattern = re.compile(r"#\d+#")

    for para in doc.paragraphs:
        detected_name = False

        # Check if the paragraph starts with a name (e.g., #10#)
        if re.match(name_pattern, para.text.strip()):
            detected_name = True

        # Alternatively, check for special formatting (bold/highlighted text) for names
        for run in para.runs:
            if run.bold or run.font.highlight_color is not None:
                detected_name = True
                break

        # If a name is detected
        if detected_name:
            if current_name:
                # Append current name and speech as a tuple
                data.append((current_name, "\n".join(current_speech)))  # Preserve line breaks
            current_name = para.text.strip('#').strip()  # Strip the # symbols
            current_speech = []  # Reset speech for the new name
        else:
            current_speech.append(para.text)  # Append paragraphs for the speech

    # Append the last name-speech pair
    if current_name:
        data.append((current_name, "\n".join(current_speech)))

    return data

# Step 2: Create a DOCX file with a table containing the extracted data
def create_output_docx(data, output_file):
    doc = Document()

    # Create a table with 2 columns: Name and Speech
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Set header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Speech'

    # Add extracted data to table
    for name, speech in data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = speech  # Speech text with original line breaks

    doc.save(output_file)

# Route for the home page
@app.route('/')
def index():
    return render_template('index.html')

# Route to handle file uploads and processing
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']
    if file.filename == '':
        return 'No selected file'

    if file and file.filename.endswith('.docx'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)

        # Process the uploaded file
        data = parse_doc(file_path)
        output_file = os.path.join(app.config['OUTPUT_FOLDER'], 'output.docx')
        create_output_docx(data, output_file)

        # Send the DOCX file back to the user
        return send_file(output_file, as_attachment=True)

    return 'Invalid file format. Please upload a .docx file.'

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
